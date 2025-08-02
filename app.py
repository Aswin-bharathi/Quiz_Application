from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, make_response
import sqlite3
import pandas as pd
from werkzeug.security import generate_password_hash, check_password_hash
import random
import io
import os
import time
from datetime import datetime
import requests
import logging
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
app.secret_key = 'intercollege_event_2025'

logging.basicConfig(level=logging.DEBUG)

def init_db():
    conn = sqlite3.connect('quiz.db')
    c = conn.cursor()
    
    c.execute('''CREATE TABLE IF NOT EXISTS students (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        lotname TEXT NOT NULL UNIQUE,
        password TEXT NOT NULL,
        quiz_status_tech TEXT DEFAULT 'not_attempted',
        quiz_status_software TEXT DEFAULT 'not_attempted'
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS questions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        quiz_type TEXT NOT NULL,
        question TEXT NOT NULL,
        option1 TEXT NOT NULL,
        option2 TEXT NOT NULL,
        option3 TEXT NOT NULL,
        option4 TEXT NOT NULL,
        answer TEXT NOT NULL
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS quiz_entries (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        entry_code TEXT NOT NULL,
        quiz_type TEXT NOT NULL
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS results (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        lotname TEXT NOT NULL,
        score INTEGER NOT NULL,
        duration INTEGER NOT NULL,
        quiz_type TEXT NOT NULL
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS admins (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        password TEXT NOT NULL
    )''')
    
    default_admin = '1'
    default_password = generate_password_hash('1')
    c.execute('INSERT OR IGNORE INTO admins (username, password) VALUES (?, ?)', 
             (default_admin, default_password))
    
    c.execute('SELECT id, lotname FROM students WHERE password IS NULL OR password = ""')
    for student in c.fetchall():
        student_id, lotname = student
        password = lotname + '@2k25'
        c.execute('UPDATE students SET password = ? WHERE id = ?', (password, student_id))
    
    conn.commit()
    conn.close()

init_db()

def admin_login_required(f):
    def wrap(*args, **kwargs):
        if 'admin' not in session or not session.get('admin'):
            flash('Please login as admin first!', 'error')
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    wrap.__name__ = f.__name__
    return wrap

@app.route('/admin_login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        conn = sqlite3.connect('quiz.db')
        c = conn.cursor()
        c.execute('SELECT * FROM admins WHERE username = ?', (username,))
        admin = c.fetchone()
        conn.close()
        
        if admin and check_password_hash(admin[2], password):
            session['admin'] = username
            return redirect(url_for('admin_dashboard'))
        flash('Invalid credentials!', 'error')
    return render_template('admin_login.html')

@app.route('/admin/dashboard')
@admin_login_required
def admin_dashboard():
    return render_template('admin_dashboard.html')

@app.route('/admin/add_team', methods=['GET', 'POST'])
@admin_login_required
def add_team():
    if request.method == 'POST':
        lotname = request.form['lotname']
        
        if not lotname:
            flash('Lot Name is required!', 'error')
            return redirect(url_for('add_team'))
            
        try:
            password = lotname + '@2k25'
            conn = sqlite3.connect('quiz.db')
            c = conn.cursor()
            c.execute('INSERT INTO students (lotname, password, quiz_status_tech, quiz_status_software) VALUES (?, ?, ?, ?)',
                     (lotname, password, 'not_attempted', 'not_attempted'))
            conn.commit()
            flash('Team added successfully!', 'success')
        except sqlite3.IntegrityError:
            flash('Team name already exists!', 'error')
        finally:
            conn.close()
            
        return redirect(url_for('show_teams'))
    return render_template('add_team.html')

@app.route('/admin/sync_teams', methods=['GET', 'POST'])
@admin_login_required
def sync_teams():
    if request.method == 'POST':
        try:
            response = requests.get('https://anjacstrata.in/mob/get_isused_lots.php')
            response.raise_for_status()
            teams_data = response.json()
            
            logging.debug(f"API Response: {teams_data}")
            
            teams_list = []
            if isinstance(teams_data, list):
                teams_list = teams_data
            elif isinstance(teams_data, dict):
                for key in ['data', 'teams', 'lots', 'records']:
                    if key in teams_data and isinstance(teams_data[key], list):
                        teams_list = teams_data[key]
                        break
                else:
                    teams_list = [teams_data]
            
            if not teams_list:
                flash('No valid team data found in response!', 'error')
                return redirect(url_for('show_teams'))
            
            conn = sqlite3.connect('quiz.db')
            c = conn.cursor()
            added = 0
            skipped = 0
            
            c.execute('SELECT lotname FROM students')
            existing_lotnames = {row[0] for row in c.fetchall()}
            
            for team in teams_list:
                try:
                    if not (isinstance(team, dict) and 'lotname' in team):
                        logging.warning(f"Invalid team data: {team}")
                        skipped += 1
                        continue
                    lotname = str(team['lotname']).strip()
                    if lotname in existing_lotnames:
                        skipped += 1
                        continue
                    password = lotname + '@2k25'
                    c.execute('INSERT INTO students (lotname, password, quiz_status_tech, quiz_status_software) VALUES (?, ?, ?, ?)',
                             (lotname, password, 'not_attempted', 'not_attempted'))
                    added += 1
                except sqlite3.IntegrityError:
                    skipped += 1
                except KeyError as e:
                    logging.warning(f"Missing field {e} in team: {team}")
                    skipped += 1
            
            conn.commit()
            conn.close()
            flash(f'Successfully added {added} teams, skipped {skipped} duplicates or invalid entries!', 'success')
        except requests.RequestException as e:
            flash(f'Error fetching data from URL: {str(e)}', 'error')
        except ValueError as e:
            flash(f'Invalid JSON data from URL: {str(e)}', 'error')
        except Exception as e:
            flash(f'Unexpected error: {str(e)}', 'error')
        
        return redirect(url_for('show_teams'))
    
    return render_template('sync_teams.html')

@app.route('/admin/update_team/<int:id>', methods=['GET', 'POST'])
@admin_login_required
def update_team(id):
    conn = sqlite3.connect('quiz.db')
    c = conn.cursor()
    
    if request.method == 'POST':
        lotname = request.form['lotname']
        
        if not lotname:
            flash('Lot Name is required!', 'error')
            conn.close()
            return redirect(url_for('update_team', id=id))
            
        try:
            password = lotname + '@2k25'
            c.execute('UPDATE students SET lotname = ?, password = ? WHERE id = ?',
                     (lotname, password, id))
            conn.commit()
            flash('Team updated successfully!', 'success')
            return redirect(url_for('show_teams'))
        except sqlite3.IntegrityError:
            flash('Team name already exists!', 'error')
        finally:
            conn.close()
    
    c.execute('SELECT * FROM students WHERE id = ?', (id,))
    team = c.fetchone()
    conn.close()
    return render_template('update_team.html', team=team)

@app.route('/admin/delete_team/<int:id>')
@admin_login_required
def delete_team(id):
    conn = sqlite3.connect('quiz.db')
    c = conn.cursor()
    c.execute('DELETE FROM students WHERE id = ?', (id,))
    conn.commit()
    conn.close()
    flash('Team deleted successfully!', 'success')
    return redirect(url_for('show_teams'))

@app.route('/admin/show_teams')
@admin_login_required
def show_teams():
    page = request.args.get('page', 1, type=int)
    per_page = 10
    search = request.args.get('search', '')
    quiz_type = request.args.get('quiz_type', 'all')
    
    conn = sqlite3.connect('quiz.db')
    c = conn.cursor()
    
    query = 'SELECT * FROM students'
    count_query = 'SELECT COUNT(*) FROM students'
    params = []
    
    if search:
        query += ' WHERE lotname LIKE ?'
        count_query += ' WHERE lotname LIKE ?'
        params.append(f'%{search}%')
    
    if quiz_type != 'all':
        status_column = 'quiz_status_tech' if quiz_type == 'Tech' else 'quiz_status_software'
        if search:
            query += f' AND {status_column} != ?'
            count_query += f' AND {status_column} != ?'
        else:
            query += f' WHERE {status_column} != ?'
            count_query += f' WHERE {status_column} != ?'
        params.append('not_attempted')
    
    query += ' ORDER BY lotname ASC'
    c.execute(count_query, params)
    total = c.fetchone()[0]
    
    c.execute(query + ' LIMIT ? OFFSET ?', params + [per_page, (page-1)*per_page])
    teams = c.fetchall()
    conn.close()
    
    total_pages = (total + per_page - 1) // per_page
    return render_template('show_teams.html', teams=teams, page=page, total_pages=total_pages, search=search, quiz_type=quiz_type)

@app.route('/admin/add_question', methods=['GET', 'POST'])
@admin_login_required
def add_question():
    if request.method == 'POST':
        quiz_type = request.form['quiz_type']
        question = request.form['question']
        option1 = request.form['option1']
        option2 = request.form['option2']
        option3 = request.form['option3']
        option4 = request.form['option4']
        answer = request.form['answer']
        
        if not all([quiz_type, question, option1, option2, option3, option4, answer]):
            flash('All fields are required!', 'error')
            return redirect(url_for('add_question'))
            
        if quiz_type not in ['Tech', 'Software']:
            flash('Invalid quiz type! Must be "Tech" or "Software".', 'error')
            return redirect(url_for('add_question'))
            
        conn = sqlite3.connect('quiz.db')
        c = conn.cursor()
        c.execute('INSERT INTO questions (quiz_type, question, option1, option2, option3, option4, answer) VALUES (?, ?, ?, ?, ?, ?, ?)',
                 (quiz_type, question, option1, option2, option3, option4, answer))
        conn.commit()
        conn.close()
        flash('Question added successfully!', 'success')
        return redirect(url_for('show_questions'))
    return render_template('add_question.html')

@app.route('/admin/delete_question/<int:id>')
@admin_login_required
def delete_question(id):
    conn = sqlite3.connect('quiz.db')
    c = conn.cursor()
    c.execute('DELETE FROM questions WHERE id = ?', (id,))
    conn.commit()
    conn.close()
    flash('Question deleted successfully!', 'success')
    return redirect(url_for('show_questions'))

@app.route('/admin/show_questions')
@admin_login_required
def show_questions():
    conn = sqlite3.connect('quiz.db')
    c = conn.cursor()
    
    quiz_type = request.args.get('quiz_type', 'all')  # Default to 'all' to show all questions
    query = 'SELECT * FROM questions'
    params = []
    
    if quiz_type != 'all':
        query += ' WHERE quiz_type = ?'
        params.append(quiz_type)
    
    c.execute(query, params)
    questions = c.fetchall()
    conn.close()
    
    return render_template('show_questions.html', questions=questions, quiz_type=quiz_type)

@app.route('/admin/upload_questions', methods=['GET', 'POST'])
@admin_login_required
def upload_questions():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file uploaded!', 'error')
            return redirect(url_for('upload_questions'))
            
        file = request.files['file']
        if file.filename == '':
            flash('No file selected!', 'error')
            return redirect(url_for('upload_questions'))
            
        if file and file.filename.endswith('.xlsx'):
            try:
                df = pd.read_excel(file)
                expected_columns = ['quiz_type', 'question', 'option1', 'option2', 'option3', 'option4', 'answer']
                if not all(col in df.columns for col in expected_columns):
                    missing_cols = [col for col in expected_columns if col not in df.columns]
                    flash(f'Missing required columns in Excel file: {", ".join(missing_cols)}', 'error')
                    return redirect(url_for('upload_questions'))
                
                conn = sqlite3.connect('quiz.db')
                c = conn.cursor()
                inserted = 0
                errors = []
                
                for index, row in df.iterrows():
                    try:
                        quiz_type = str(row['quiz_type']).strip()
                        if quiz_type not in ['Tech', 'Software']:
                            errors.append(f"Row {index + 2}: Invalid quiz_type '{quiz_type}'. Must be 'Tech' or 'Software'.")
                            continue
                        
                        question = str(row['question']).strip()
                        option1 = str(row['option1']).strip()
                        option2 = str(row['option2']).strip()
                        option3 = str(row['option3']).strip()
                        option4 = str(row['option4']).strip()
                        answer = str(row['answer']).strip()
                        
                        if not all([question, option1, option2, option3, option4, answer]):
                            errors.append(f"Row {index + 2}: Missing required fields.")
                            continue
                        
                        c.execute('INSERT INTO questions (quiz_type, question, option1, option2, option3, option4, answer) VALUES (?, ?, ?, ?, ?, ?, ?)',
                                 (quiz_type, question, option1, option2, option3, option4, answer))
                        inserted += 1
                    except Exception as e:
                        errors.append(f"Row {index + 2}: Error inserting - {str(e)}")
                        logging.error(f"Error inserting row {index + 2}: {row.to_dict()} - {str(e)}")
                
                conn.commit()
                conn.close()
                
                if inserted > 0:
                    flash(f'Successfully uploaded {inserted} questions!', 'success')
                else:
                    flash('No questions were uploaded due to errors!', 'error')
                
                if errors:
                    flash('Errors encountered during upload: ' + '; '.join(errors), 'error')
                    
            except Exception as e:
                flash(f'Error reading Excel file: {str(e)}', 'error')
        else:
            flash('Please upload a valid .xlsx file!', 'error')
            
        return redirect(url_for('show_questions'))
    return render_template('upload_questions.html')

@app.route('/admin/set_entry_code', methods=['GET', 'POST'])
@admin_login_required
def set_entry_code():
    if request.method == 'POST':
        entry_code = request.form['entry_code']
        quiz_type = request.form['quiz_type']
        if not all([entry_code, quiz_type]):
            flash('Entry code and quiz type are required!', 'error')
            return redirect(url_for('set_entry_code'))
        
        if quiz_type not in ['Tech', 'Software']:
            flash('Invalid quiz type! Must be "Tech" or "Software".', 'error')
            return redirect(url_for('set_entry_code'))
            
        conn = sqlite3.connect('quiz.db')
        c = conn.cursor()
        c.execute('INSERT INTO quiz_entries (entry_code, quiz_type) VALUES (?, ?)', (entry_code, quiz_type))
        conn.commit()
        conn.close()
        flash('Entry code set successfully!', 'success')
        return redirect(url_for('admin_dashboard'))
    return render_template('set_entry_code.html')

def format_duration(seconds):
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"

@app.route('/admin/view_results')
@admin_login_required
def view_results():
    conn = sqlite3.connect('quiz.db')
    c = conn.cursor()
    
    page = request.args.get('page', 1, type=int)
    per_page = 10
    search = request.args.get('search', '').strip()
    quiz_type = request.args.get('quiz_type', 'all').strip()
    
    query = '''
        SELECT results.*, students.lotname 
        FROM results 
        JOIN students ON results.lotname = students.lotname
    '''
    count_query = '''
        SELECT COUNT(*) 
        FROM results 
        JOIN students ON results.lotname = students.lotname
    '''
    params = []
    
    if quiz_type != 'all':
        query += ' WHERE results.quiz_type = ?'
        count_query += ' WHERE results.quiz_type = ?'
        params.append(quiz_type)
        print(f"Added quiz_type filter: {quiz_type}")
    
    if search:
        condition = ' AND ' if quiz_type != 'all' else ' WHERE '
        query += f'{condition}(students.lotname LIKE ?)'
        count_query += f'{condition}(students.lotname LIKE ?)'
        params.append(f'%{search}%')
        print(f"Added search filter: {search}")
    
    query += ' ORDER BY results.score DESC, results.duration ASC, results.lotname ASC'
    print(f"Count Query: {count_query}")
    print(f"Params: {params}")
    c.execute(count_query, params)
    total = c.fetchone()[0]
    
    c.execute(query + ' LIMIT ? OFFSET ?', params + [per_page, (page-1)*per_page])
    results = c.fetchall()
    
    # Modified top_teams query to include WHERE clause when quiz_type != 'all'
    top_teams_query = '''
        SELECT students.lotname, results.score, results.duration, results.quiz_type 
        FROM results 
        JOIN students ON results.lotname = students.lotname
    '''
    top_teams_params = []
    if quiz_type != 'all':
        top_teams_query += ' WHERE results.quiz_type = ?'
        top_teams_params.append(quiz_type)
    top_teams_query += ' ORDER BY results.score DESC, results.duration ASC LIMIT 5'
    print(f"Top Teams Query: {top_teams_query}")
    print(f"Top Teams Params: {top_teams_params}")
    c.execute(top_teams_query, top_teams_params)
    top_teams = c.fetchall()
    conn.close()
    
    total_pages = (total + per_page - 1) // per_page
    formatted_results = [(r[0], r[1], r[2], format_duration(r[3]), r[4], r[5]) for r in results]
    formatted_top_teams = [(t[0], t[1], format_duration(t[2]), t[3]) for t in top_teams]
    
    if request.args.get('format') == 'word':
        doc = Document()
        doc.add_heading(f'Top 5 Teams - {"All" if quiz_type == "all" else quiz_type} Quiz', 0)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Team Name'
        hdr_cells[1].text = 'Score'
        hdr_cells[2].text = 'Duration (HH:MM:SS)'
        for lotname, score, duration, _ in formatted_top_teams:
            row_cells = table.add_row().cells
            row_cells[0].text = lotname
            row_cells[1].text = str(score)
            row_cells[2].text = duration
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        response = make_response(output.getvalue())
        response.headers['Content-Disposition'] = f'attachment; filename=top_5_teams_{"all" if quiz_type == "all" else quiz_type.lower()}.docx'
        response.headers['Content-type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response
    
    if request.args.get('format') == 'excel':
        df = pd.DataFrame([(r[5], r[2], r[3]) for r in formatted_results], columns=['lotname', 'score', 'duration'])
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Results')
        output.seek(0)
        response = make_response(output.getvalue())
        response.headers['Content-Disposition'] = f'attachment; filename=results_{"all" if quiz_type == "all" else quiz_type.lower()}.xlsx'
        response.headers['Content-type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        return response
    
    return render_template('view_results.html', results=formatted_results, top_teams=formatted_top_teams, page=page, total_pages=total_pages, search=search, quiz_type=quiz_type)

@app.route('/logout')
@admin_login_required
def admin_logout():
    session.pop('admin', None)
    flash('Logged out successfully!', 'success')
    return redirect(url_for('admin_login'))

@app.route('/quiz_login', methods=['GET', 'POST'])
def quiz_login():
    if request.method == 'POST':
        lotname = request.form['lotname']
        password = request.form['password']
        entry_code = request.form['entry_code']
        
        conn = sqlite3.connect('quiz.db')
        c = conn.cursor()
        
        c.execute('SELECT quiz_type FROM quiz_entries WHERE entry_code = ?', (entry_code,))
        entry = c.fetchone()
        
        if not entry:
            flash('Invalid Entry Code!', 'error')
            conn.close()
            return render_template('quiz_login.html')
        
        quiz_type = entry[0]
        quiz_status_column = f'quiz_status_{quiz_type.lower()}'
        
        c.execute(f'SELECT * FROM students WHERE lotname = ? AND password = ? AND {quiz_status_column} = ?',
                 (lotname, password, 'not_attempted'))
        team = c.fetchone()
        
        if team:
            session['team'] = {'lotname': team[1], 'quiz_type': quiz_type}
            return redirect(url_for('quiz_start'))
        else:
            flash('Invalid Lot Name, Password, or quiz already attempted!', 'error')
        
        conn.close()
    return render_template('quiz_login.html')

@app.route('/quiz_start')
def quiz_start():
    if 'team' not in session:
        flash('Please login first!', 'error')
        return redirect(url_for('quiz_login'))
        
    conn = sqlite3.connect('quiz.db')
    c = conn.cursor()
    c.execute('SELECT question, option1, option2, option3, option4, answer FROM questions WHERE quiz_type = ?', (session['team']['quiz_type'],))
    questions = c.fetchall()
    conn.close()
    
    if not questions:
        flash('No questions available for this quiz type!', 'error')
        return redirect(url_for('quiz_login'))
    
    random.shuffle(questions)
    session['questions'] = questions[:10]
    session['tab_switches'] = 0
    session['quiz_active'] = True
    session['start_time'] = time.time()
    
    return render_template('quiz_start.html', questions=session['questions'])

@app.route('/quiz_submit', methods=['POST'])
def quiz_submit():
    if 'team' not in session:
        return jsonify({'error': 'Not logged in'}), 403
        
    answers = request.form.to_dict()
    questions = session['questions']
    score = 0
    
    for index, q in enumerate(questions, 1):  # Start index from 1 to match loop.index
        q_id = f"q{index}"
        answer_key = f"q{index}-answer"
        if q_id in answers and answer_key in answers and answers[q_id] == answers[answer_key]:
            score += 1
    
    duration = int(time.time() - session['start_time'])
    quiz_type = session['team']['quiz_type']
    
    conn = sqlite3.connect('quiz.db')
    c = conn.cursor()
    c.execute('INSERT INTO results (lotname, score, duration, quiz_type) VALUES (?, ?, ?, ?)',
             (session['team']['lotname'], score, duration, quiz_type))
    c.execute(f'UPDATE students SET quiz_status_{quiz_type.lower()} = ? WHERE lotname = ?',
             ('attempted', session['team']['lotname']))
    conn.commit()
    conn.close()
    
    session.clear()
    session['quiz_completed'] = True
    
    return redirect(url_for('quiz_result'))

@app.route('/quiz_result')
def quiz_result():
    if not session.get('quiz_completed', False):
        flash('Please complete the quiz first!', 'error')
        return redirect(url_for('quiz_login'))
    # Disable browser cache to prevent back navigation to cached pages
    response = make_response(render_template('quiz_result.html'))
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

@app.route('/quiz/track_tab_switch', methods=['POST'])
def track_tab_switch():
    if 'team' in session and session.get('quiz_active', False):
        session['tab_switches'] = session.get('tab_switches', 0) + 1
        if session['tab_switches'] >= 1:
            session['quiz_active'] = False
            return jsonify({'auto_submit': True})
    return jsonify({'auto_submit': False})

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=True, port=port, host="0.0.0.0")